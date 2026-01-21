from docx import Document
import json
import re

FILES = [
    "soal siberr.docx",
    "soal siber .docx",
    "QuizM10 Unified.docx",
    "Quiz M09 Unified.docx",
    "quiz cyber1.docx",
    "sesuaikan lagi pertanyaan.docx"
]

questions = []
qid = 1

def is_blue(run):
    if run.font.color and run.font.color.rgb:
        return str(run.font.color.rgb) == "0000FF"
    return False

for file in FILES:
    doc = Document(file)
    i = 0

    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        text = p.text.strip()

        if re.match(r"^\d+\.|^Soal|^Question", text):
            q_text = re.sub(r"^\d+\.\s*", "", text)
            i += 1

            options = {}
            correct = []

            while i < len(doc.paragraphs):
                line = doc.paragraphs[i]
                raw = line.text.strip()

                if re.match(r"^[A-Ea-e]\.", raw):
                    key = raw[0].upper()
                    val = raw[2:].strip()
                    options[key] = val

                    for run in line.runs:
                        if "✅" in run.text or run.bold or is_blue(run):
                            correct.append(key)

                    i += 1
                else:
                    break

            q_type = "multiple" if len(correct) > 1 else "single"

            if options and correct:
                questions.append({
                    "id": qid,
                    "question": q_text,
                    "options": options,
                    "correct": list(set(correct)),
                    "type": q_type,
                    "weight": 1
                })
                qid += 1
        else:
            i += 1

with open("questions.json", "w", encoding="utf-8") as f:
    json.dump(questions, f, indent=2, ensure_ascii=False)

print("Total soal valid:", len(questions))
